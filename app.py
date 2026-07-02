import streamlit as st
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
import pandas as pd
import os
import sys
import logging
import re
from datetime import datetime
import io
import tempfile

# ============================================
# 1. VERSION & INFO
# ============================================

__version__ = "2.0.0"
__author__ = "BP2JK Kalteng"
__last_modified__ = "2026-07-02"

def get_app_info():
    """Mendapatkan informasi aplikasi"""
    return {
        'name': 'APSO Ultimate',
        'version': __version__,
        'author': __author__,
        'last_modified': __last_modified__,
        'tnde_compliant': True
    }

# ============================================
# 2. KONFIGURASI & LOGGING
# ============================================

def get_default_paths():
    """Mendapatkan path default dengan fallback cross-platform"""
    if os.name == 'nt':  # Windows
        default_base = r"C:\doklib"
    else:  # Linux/Mac
        default_base = os.path.expanduser("~/doklib")
    
    return {
        'base': default_base,
        'templates': os.path.join(default_base, 'tnd'),
        'output': default_base,
        'logs': os.path.join(default_base, 'logs')
    }

def setup_logging():
    """Setup logging untuk debugging dan audit trail"""
    paths = get_default_paths()
    log_dir = paths['logs']
    os.makedirs(log_dir, exist_ok=True)
    
    log_file = os.path.join(log_dir, f'apso_{datetime.now().strftime("%Y%m%d")}.log')
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
    return logging.getLogger(__name__)

logger = setup_logging()

# ============================================
# 2. STANDAR TNDE (Tata Naskah Dinas Elektronik)
# ============================================

class TNDEStandard:
    """Standar penulisan surat sesuai TNDE"""
    
    MARGIN_LEFT = 4.0
    MARGIN_RIGHT = 3.0
    MARGIN_TOP = 4.0
    MARGIN_BOTTOM = 3.0
    FONT_NAME = 'Times New Roman'
    FONT_SIZE = 12
    LINE_SPACING = 1.5
    
    @staticmethod
    def format_nomor_surat(no_urut, bulan, tahun, kode_unit='BP2JK-KT'):
        return f"{no_urut:03d}/{kode_unit}/{bulan:02d}/{tahun}"
    
    @staticmethod
    def apply_standards(doc):
        for section in doc.sections:
            section.left_margin = Cm(TNDEStandard.MARGIN_LEFT)
            section.right_margin = Cm(TNDEStandard.MARGIN_RIGHT)
            section.top_margin = Cm(TNDEStandard.MARGIN_TOP)
            section.bottom_margin = Cm(TNDEStandard.MARGIN_BOTTOM)
        
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Normal') or paragraph.style.name == 'Normal':
                for run in paragraph.runs:
                    run.font.name = TNDEStandard.FONT_NAME
                    run.font.size = Pt(TNDEStandard.FONT_SIZE)
                paragraph.paragraph_format.line_spacing = TNDEStandard.LINE_SPACING
                paragraph.paragraph_format.space_after = Pt(0)
                if paragraph.alignment is None:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    @staticmethod
    def validate_nomor_surat(nomor_surat):
        pattern = r'^\d{3}/[A-Z0-9-]+/\d{2}/\d{4}$'
        return bool(re.match(pattern, nomor_surat))

# ============================================
# 4. FUNGSI REPLACE TEXT (OPTIMASI SPLIT-RUN)
# ============================================

def merge_paragraph_runs(paragraph):
    """
    Menyatukan runs yang terpecah (split-runs) HANYA jika memiliki 
    format yang sama persis. Menjaga teks bold/italic lain di paragraf tetap aman.
    """
    if len(paragraph.runs) <= 1:
        return

    i = 0
    while i < len(paragraph.runs) - 1:
        r1 = paragraph.runs[i]
        r2 = paragraph.runs[i + 1]

        # Validasi ketat: Gabungkan HANYA jika semua properti formatnya sama persis
        if (r1.bold == r2.bold and
            r1.italic == r2.italic and
            r1.underline == r2.underline and
            getattr(r1.font, 'name', None) == getattr(r2.font, 'name', None) and
            getattr(r1.font, 'size', None) == getattr(r2.font, 'size', None)):
            
            # Gabungkan teks run kedua ke run pertama
            r1.text += r2.text
            # Hapus run kedua dari XML Word secara bersih
            r2._element.getparent().remove(r2._element)
            # Jangan naikkan indeks i, karena r1 yang baru akan dicek lagi dengan run berikutnya
        else:
            i += 1

def smart_replace_text_with_placeholders(paragraph, data_teks):
    """
    Mengganti placeholder dengan mempertahankan format dokumen asli.
    Sinkron 100% dengan skrip utama dan mempertahankan logika fallback key.
    """
    # 1. Jalankan penggabungan run yang aman (tidak destruktif)
    merge_paragraph_runs(paragraph)
    
    if not paragraph.runs:
        return
        
    full_text = paragraph.text
    
    # 2. Cek awal apakah ada placeholder yang cocok
    butuh_replace = False
    for key in data_teks.keys():
        if key in full_text:
            butuh_replace = True
            break
            
    if not butuh_replace:
        return

    # 3. Lakukan penggantian langsung pada run yang sudah bersih
    for run in paragraph.runs:
        for key, value in data_teks.items():
            # Penggantian utama (Exact Match)
            if key in run.text:
                run.text = run.text.replace(key, str(value))
                
            # Logika Antisipasi jika key ditulis tanpa bracket di template
            clean_key = key.strip("[]{}")
            if clean_key == key:
                continue
            if clean_key in run.text and f"[{clean_key}]" not in full_text:
                run.text = run.text.replace(clean_key, str(value))

def replace_alinea_shortcuts(paragraph, data_teks):
    """
    Fungsi terpisah untuk mengganti shortcut alinea
    Karena shortcut alinea tidak menggunakan placeholder pattern
    """
    merge_paragraph_runs(paragraph)
    
    text = paragraph.text
    
    replacements = {
        "(Alinea pembuka)": data_teks.get("pembuka", ""),
        "(Alinea isi)": data_teks.get("isi", ""),
        "(Alinea penutup)": data_teks.get("penutup", ""),
        "(Alinea pembuka dan alinea isi)": data_teks.get("pembuka_isi", ""),
    }
    
    for key, value in replacements.items():
        if key in text:
            if value:
                paragraph.clear()
                paragraph.add_run(value)
            else:
                paragraph.text = ""
            return True
    
    return False

def process_all_placeholders(doc, data_teks):
    """
    Proses semua placeholder di seluruh dokumen
    OPTIMASI: Batch processing
    """
    # Kumpulkan semua paragraf
    paragraphs_to_process = []
    
    # Paragraf biasa
    for paragraph in doc.paragraphs:
        paragraphs_to_process.append(paragraph)
    
    # Paragraf dalam tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs_to_process.append(paragraph)
    
    # Proses semua paragraf
    for paragraph in paragraphs_to_process:
        if not replace_alinea_shortcuts(paragraph, data_teks):
            smart_replace_text_with_placeholders(paragraph, data_teks)

def find_matching_table(doc, cols_web):
    """Cari tabel yang cocok dengan kolom web"""
    for table in doc.tables:
        if len(table.rows) > 0:
            cols_word = [cell.text.strip().lower() for cell in table.rows[0].cells]
            matches = sum(1 for c in cols_web if c in cols_word)
            if matches >= 2 or (len(cols_web) == 1 and cols_web[0] in cols_word):
                return table, cols_word
    return None, None

# ============================================
# 4. FUNGSI PREVIEW DOKUMEN (DIPERBAIKI)
# ============================================

def generate_preview_document(template_path, data_teks, data_tabel, file_gambar, penandatangan_list):
    """
    Generate preview dokumen dengan semua penggantian
    """
    try:
        doc = Document(template_path)
        logger.info("Preview: Template berhasil dibuka")
        
        # 1. Proses semua placeholder di seluruh dokumen
        process_all_placeholders(doc, data_teks)
        logger.info("Preview: Placeholder berhasil diganti")
        
        # 2. Proses tabel dinamis (preview hanya 3 baris pertama)
        if data_tabel is not None and not data_tabel.empty:
            cols_web = [c.lower().strip() for c in data_tabel.columns]
            table, cols_word = find_matching_table(doc, cols_web)
            
            if table:
                logger.info("Preview: Menemukan tabel yang cocok di Word")
                preview_data = data_tabel.head(3)
                for _, row_data in preview_data.iterrows():
                    new_row = table.add_row()
                    for col_name in data_tabel.columns:
                        c_lower = col_name.lower().strip()
                        if c_lower in cols_word:
                            target_idx = cols_word.index(c_lower)
                            value = "" if pd.isna(row_data[col_name]) else str(row_data[col_name])
                            new_row.cells[target_idx].text = value
                
                if len(data_tabel) > 3:
                    doc.add_paragraph(f"... (dan {len(data_tabel) - 3} baris data lainnya untuk preview)")
            else:
                logger.info("Preview: Tidak ada tabel yang cocok, membuat tabel lampiran")
                doc.add_paragraph("\nLampiran Data Tabel (Preview):")
                tabel_word = doc.add_table(rows=1, cols=len(data_tabel.columns))
                tabel_word.style = 'Table Grid'
                
                for i, col_name in enumerate(data_tabel.columns):
                    tabel_word.rows[0].cells[i].text = str(col_name)
                
                preview_data = data_tabel.head(3)
                for _, row in preview_data.iterrows():
                    row_cells = tabel_word.add_row().cells
                    for i, item in enumerate(row):
                        row_cells[i].text = "" if pd.isna(item) else str(item)
                
                if len(data_tabel) > 3:
                    doc.add_paragraph(f"... (dan {len(data_tabel) - 3} baris data lainnya)")
        
        # 3. Proses gambar (thumbnail untuk preview)
        if file_gambar is not None:
            logger.info("Preview: Memproses gambar")
            doc.add_paragraph("\nLampiran Gambar / Dokumentasi (Preview):")
            p_gambar = doc.add_paragraph()
            r_gambar = p_gambar.add_run()
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                tmp_file.write(file_gambar.getbuffer())
                temp_img_path = tmp_file.name
            
            try:
                r_gambar.add_picture(temp_img_path, width=Inches(2.5))
            finally:
                os.unlink(temp_img_path)
        
        # 4. Proses penandatangan (preview)
        for paragraph in doc.paragraphs:
            if "Nama Lengkap Tanpa Gelar" in paragraph.text:
                paragraph.text = ""
        
        if penandatangan_list:
            logger.info(f"Preview: Memproses {len(penandatangan_list)} penandatangan")
            doc.add_paragraph("\n" * 2)
            tempat_tanggal = data_teks.get("tempat_tanggal", f"Palangka Raya, {datetime.now().strftime('%d %B YYYY')}")
            doc.add_paragraph(tempat_tanggal)
            doc.add_paragraph("\n" * 2)
            
            num_ttd = len(penandatangan_list)
            tabel_ttd = doc.add_table(rows=2, cols=num_ttd)
            tabel_ttd.autofit = True
            
            for idx, ptd in enumerate(penandatangan_list):
                cell_jabatan = tabel_ttd.rows[0].cells[idx]
                cell_jabatan.text = f"{ptd['jabatan']}\n\n\n\n"
                cell_jabatan.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                cell_nama = tabel_ttd.rows[1].cells[idx]
                cell_nama.text = f"({ptd['nama']})\nNIP. {ptd['nip']}"
                cell_nama.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 5. Tambahkan watermark preview
        doc.add_paragraph("\n" + "="*50)
        doc.add_paragraph("PREVIEW DOKUMEN - BUKAN VERSI FINAL")
        doc.add_paragraph("="*50)
        
        # Simpan ke buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("Preview: Dokumen preview berhasil dibuat")
        return buffer
        
    except Exception as e:
        logger.error(f"Error generating preview: {e}")
        raise
def extract_text_preview(doc_buffer, max_length=3000):
    """Ekstrak teks dari dokumen untuk preview"""
    try:
        doc = Document(doc_buffer)
        text_parts = []
        total_length = 0
        
        # Ambil teks dari paragraf
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                total_length += len(paragraph.text)
                if total_length > max_length:
                    text_parts.append("\n... (preview terpotong)")
                    break
        
        # Tambahkan teks dari tabel jika masih ada ruang
        if total_length < max_length:
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
                        total_length += len(row_text)
                        if total_length > max_length:
                            text_parts.append("\n... (preview terpotong)")
                            break
                if total_length > max_length:
                    break
        
        return "\n".join(text_parts)
    except Exception as e:
        return f"Error extracting text: {e}"

def show_preview_dialog(preview_buffer, preview_data, template_name):
    """
    Tampilkan preview dokumen dalam dialog
    """
    st.markdown("---")
    st.subheader("👁️ Preview Dokumen")
    
    # Informasi preview
    st.info(f"📄 **Preview untuk template:** {template_name}")
    st.caption("⚠️ Preview ini adalah versi ringkas dengan watermark. File final akan lengkap.")
    
    # Tabs untuk berbagai mode preview
    preview_tab1, preview_tab2, preview_tab3 = st.tabs([
        "📄 Download Preview", "📝 Preview Teks", "📊 Ringkasan Data"
    ])
    
    with preview_tab1:
        st.success("📄 **Download file preview** untuk melihat tampilan lengkap di Microsoft Word")
        
        # Download preview
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="📥 Download Preview (.docx)",
                data=preview_buffer,
                file_name=f"preview_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col2:
            # Opsi untuk membuka di Word (hanya Windows)
            if os.name == 'nt':
                if st.button("📂 Buka di Microsoft Word", use_container_width=True):
                    try:
                        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                            tmp.write(preview_buffer.getvalue())
                            tmp_path = tmp.name
                        os.startfile(tmp_path)
                        st.success(f"✅ Dokumen preview dibuka di: {tmp_path}")
                    except Exception as e:
                        st.error(f"Gagal membuka dokumen: {e}")
            else:
                st.info("💡 Fitur 'Buka di Word' hanya tersedia di Windows")
        
        # Tampilkan ukuran file
        file_size = len(preview_buffer.getvalue()) / 1024
        st.caption(f"📊 Ukuran file preview: {file_size:.1f} KB")
    
    with preview_tab2:
        st.subheader("📝 Preview Teks")
        
        try:
            # Reset buffer position
            preview_buffer.seek(0)
            preview_text = extract_text_preview(preview_buffer)
            
            # Tampilkan dalam text area dengan scroll
            st.text_area(
                "Isi Dokumen (Preview Teks)",
                preview_text,
                height=400,
                key="preview_text_area",
                help="Scroll untuk melihat seluruh teks"
            )
            
            # Statistik teks
            preview_buffer.seek(0)
            doc = Document(preview_buffer)
            st.caption(f"📊 Total karakter: {len(preview_text)} | Paragraf: {len(doc.paragraphs)} | Tabel: {len(doc.tables)}")
            
        except Exception as e:
            st.error(f"Gagal menampilkan preview teks: {e}")
    
    with preview_tab3:
        st.subheader("📊 Ringkasan Data yang Akan Diproses")
        
        if preview_data:
            col_sum1, col_sum2, col_sum3 = st.columns(3)
            
            with col_sum1:
                st.metric("📝 Template", preview_data.get('template', '-'), help="Template surat yang digunakan")
                st.metric("📄 Nomor Surat", preview_data.get('nomor_surat', '-'), help="Nomor surat yang akan dicetak")
            
            with col_sum2:
                st.metric("👥 Penandatangan", preview_data.get('jumlah_ttd', 0), help="Jumlah penandatangan")
                st.metric("📊 Tabel Data", f"{preview_data.get('baris_tabel', 0)} baris", help="Jumlah baris data tabel")
            
            with col_sum3:
                st.metric("🖼️ Gambar", "✅ Ada" if preview_data.get('ada_gambar') else "❌ Tidak", help="Ada lampiran gambar?")
                st.metric("📝 Alinea", preview_data.get('alinea_count', 0), help="Jumlah alinea yang diisi")
            
            # Detail tambahan
            with st.expander("📋 Detail Lengkap Data", expanded=False):
                st.json(preview_data)
    
    # Tombol aksi
    st.markdown("---")
    st.subheader("⚡ Aksi Selanjutnya")
    
    col_confirm1, col_confirm2, col_confirm3, col_confirm4 = st.columns([1, 1, 1, 1])
    
    with col_confirm1:
        if st.button("✅ Cetak Sekarang", type="primary", use_container_width=True):
            st.session_state['confirm_print'] = True
            st.session_state['show_preview'] = False
            st.rerun()
    
    with col_confirm2:
        if st.button("✏️ Edit Data", use_container_width=True):
            st.session_state['show_preview'] = False
            st.rerun()
    
    with col_confirm3:
        if st.button("📋 Copy Data", use_container_width=True):
            st.info("Data telah disalin ke clipboard (simulasi)")
    
    with col_confirm4:
        if st.button("❌ Tutup Preview", use_container_width=True):
            st.session_state['show_preview'] = False
            st.rerun()

# ============================================
# 5. FUNGSI PROSES DOKUMEN UTAMA (DIPERBAIKI)
# ============================================

def proses_dokumen_universal(template_path, output_path, data_teks, data_tabel, 
                             file_gambar, penandatangan_list, metadata=None):
    """
    Proses dokumen final dengan standar TNDE
    PERBAIKAN:
    - Menggunakan fungsi process_all_placeholders yang sudah dioptimasi
    - Mapping tabel berdasarkan nama kolom
    """
    logger.info(f"Mulai memproses dokumen: {template_path}")
    
    try:
        doc = Document(template_path)
        logger.info("Template berhasil dibuka")
        
        # 1. Terapkan standar TNDE
        TNDEStandard.apply_standards(doc)
        
        # 2. Proses semua placeholder di seluruh dokumen
        process_all_placeholders(doc, data_teks)
        logger.info("Placeholder berhasil diganti")
        
        # 3. Handling tabel dinamis (full data)
        if data_tabel is not None and not data_tabel.empty:
            tabel_bawaan_terisi = False
            cols_web = [c.lower().strip() for c in data_tabel.columns]
            
            for table in doc.tables:
                if len(table.rows) > 0:
                    cols_word = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    jumlah_cocok = sum(1 for c in cols_web if c in cols_word)
                    
                    if jumlah_cocok >= 2 or (len(cols_web) == 1 and cols_web[0] in cols_word):
                        logger.info(f"Menemukan tabel yang cocok di Word, mengisi {len(data_tabel)} baris data")
                        
                        for index, row_data in data_tabel.iterrows():
                            new_row = table.add_row()
                            
                            # ISI DATA BERDASARKAN COCOK NAMA KOLOM
                            for col_name in data_tabel.columns:
                                c_lower = col_name.lower().strip()
                                if c_lower in cols_word:
                                    target_idx = cols_word.index(c_lower)
                                    value = "" if pd.isna(row_data[col_name]) else str(row_data[col_name])
                                    new_row.cells[target_idx].text = value
                        
                        tabel_bawaan_terisi = True
                        break
            
            if not tabel_bawaan_terisi:
                logger.info("Tidak ada tabel yang cocok, membuat tabel lampiran")
                
                doc.add_paragraph("\nLampiran Data Tabel:")
                tabel_word = doc.add_table(rows=1, cols=len(data_tabel.columns))
                tabel_word.style = 'Table Grid'
                
                hdr_cells = tabel_word.rows[0].cells
                for i, col_name in enumerate(data_tabel.columns):
                    hdr_cells[i].text = str(col_name)
                
                for index, row in data_tabel.iterrows():
                    row_cells = tabel_word.add_row().cells
                    for i, item in enumerate(row):
                        row_cells[i].text = "" if pd.isna(item) else str(item)
        
        # 4. Handling gambar
        if file_gambar is not None:
            logger.info("Memproses gambar lampiran")
            
            doc.add_paragraph("\nLampiran Gambar / Dokumentasi:")
            p_gambar = doc.add_paragraph()
            r_gambar = p_gambar.add_run()
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                tmp_file.write(file_gambar.getbuffer())
                temp_img_path = tmp_file.name
            
            try:
                r_gambar.add_picture(temp_img_path, width=Inches(4.0))
                logger.info("Gambar berhasil ditambahkan")
            finally:
                os.unlink(temp_img_path)
        
        # 5. Handling signature
        for paragraph in doc.paragraphs:
            if "Nama Lengkap Tanpa Gelar" in paragraph.text:
                paragraph.text = ""
        
        if penandatangan_list:
            logger.info(f"Memproses {len(penandatangan_list)} penandatangan")
            
            doc.add_paragraph("\n" * 2)
            tempat_tanggal = data_teks.get("tempat_tanggal", f"Palangka Raya, {datetime.now().strftime('%d %B %Y')}")
            doc.add_paragraph(tempat_tanggal)
            doc.add_paragraph("\n" * 2)
            
            num_ttd = len(penandatangan_list)
            tabel_ttd = doc.add_table(rows=2, cols=num_ttd)
            tabel_ttd.autofit = True
            
            for idx, ptd in enumerate(penandatangan_list):
                cell_jabatan = tabel_ttd.rows[0].cells[idx]
                cell_jabatan.text = f"{ptd['jabatan']}\n\n\n\n"
                cell_jabatan.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                cell_nama = tabel_ttd.rows[1].cells[idx]
                cell_nama.text = f"({ptd['nama']})\nNIP. {ptd['nip']}"
                cell_nama.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 6. Metadata
        if metadata:
            core_props = doc.core_properties
            if metadata.get('judul'):
                core_props.title = metadata['judul']
            if metadata.get('perihal'):
                core_props.subject = metadata['perihal']
            if metadata.get('pembuat'):
                core_props.author = metadata['pembuat']
            core_props.created = datetime.now()
            core_props.keywords = 'TNDE, BP2JK Kalteng, Surat Dinas'
        
        # 7. Simpan
        doc.save(output_path)
        logger.info(f"Dokumen berhasil disimpan: {output_path}")
        
        return True
        
    except Exception as e:
        logger.error(f"Error dalam proses dokumen: {str(e)}")
        raise

# ============================================
# 6. FUNGSI VALIDASI
# ============================================

def validate_inputs(data_teks, data_tabel, penandatangan_list):
    """Validasi input sebelum diproses"""
    errors = []
    warnings = []
    
    nomor = data_teks.get('[NOMOR_SURAT]', '')
    if not nomor:
        errors.append("❌ Nomor surat wajib diisi")
    elif not TNDEStandard.validate_nomor_surat(nomor):
        warnings.append(f"⚠️ Format nomor surat '{nomor}' tidak sesuai standar TNDE (format: 001/BP2JK-KT/01/2026)")
    
    if not penandatangan_list or len(penandatangan_list) == 0:
        errors.append("❌ Minimal 1 penandatangan harus diisi")
    
    for idx, ptd in enumerate(penandatangan_list):
        if not ptd.get('nama', '').strip():
            errors.append(f"❌ Nama penandatangan ke-{idx+1} wajib diisi")
        if not ptd.get('nip', '').strip():
            errors.append(f"❌ NIP penandatangan ke-{idx+1} wajib diisi")
        if not ptd.get('jabatan', '').strip():
            warnings.append(f"⚠️ Jabatan penandatangan ke-{idx+1} belum diisi")
    
    if data_tabel is not None and not data_tabel.empty:
        if any(col.strip() == '' for col in data_tabel.columns):
            errors.append("❌ Nama kolom tabel tidak boleh kosong")
        
        if data_tabel.isnull().all().all():
            warnings.append("⚠️ Tabel tidak berisi data, hanya kolom yang terisi")
    
    return errors, warnings

# ============================================
# 7. MAIN STREAMLIT APP
# ============================================

def main():
    # Konfigurasi page
    st.set_page_config(
        page_title="APSO Ultimate - BP2JK Kalteng", 
        layout="wide", 
        page_icon="📝",
        initial_sidebar_state="collapsed"
    )
    
    # Inisialisasi session state
    if 'show_preview' not in st.session_state:
        st.session_state['show_preview'] = False
    if 'confirm_print' not in st.session_state:
        st.session_state['confirm_print'] = False
    if 'preview_buffer' not in st.session_state:
        st.session_state['preview_buffer'] = None
    if 'preview_data' not in st.session_state:
        st.session_state['preview_data'] = None
    if 'gambar_key' not in st.session_state:
        st.session_state['gambar_key'] = "uploader_awal"
    
    # CSS custom
    st.markdown("""
    <style>
        .main-header {
            background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
            padding: 1.5rem;
            border-radius: 10px;
            color: white;
            margin-bottom: 2rem;
        }
        .main-header h1 {
            color: white;
            margin: 0;
            font-size: 2rem;
        }
        .main-header p {
            color: #e0e0e0;
            margin: 0;
        }
        .main-header .subtitle {
            color: #b0b0b0;
            font-size: 0.9rem;
        }
        .info-box {
            background: #f0f2f6;
            padding: 1rem;
            border-radius: 10px;
            border-left: 5px solid #2a5298;
            margin: 1rem 0;
        }
        .warning-box {
            background: #fff3cd;
            padding: 1rem;
            border-radius: 10px;
            border-left: 5px solid #ffc107;
            margin: 1rem 0;
        }
        .preview-badge {
            background: #ff6b6b;
            color: white;
            padding: 0.25rem 1rem;
            border-radius: 20px;
            font-size: 0.8rem;
            font-weight: bold;
            display: inline-block;
        }
        .stButton button {
            border-radius: 8px;
            font-weight: 500;
        }
        .stButton button[kind="primary"] {
            background: #2a5298;
        }
        .stButton button:hover {
            transform: translateY(-2px);
            transition: all 0.3s ease;
        }
        .optimization-badge {
            background: #28a745;
            color: white;
            padding: 0.25rem 1rem;
            border-radius: 20px;
            font-size: 0.8rem;
            font-weight: bold;
            display: inline-block;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🚀 APSO Ultimate</h1>
        <p>Aplikasi Surat Otomatis - BP2JK Wilayah Kalimantan Tengah</p>
        <p class="subtitle">Sesuai Standar TNDE (Tata Naskah Dinas Elektronik)</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Badge optimasi
    st.markdown('<span class="optimization-badge">⚡ OPTIMASI SPLIT-RUN</span>', unsafe_allow_html=True)
    st.caption("✅ Kode telah dioptimasi untuk menangani placeholder yang terpecah di XML Word")
    
    # Tampilkan badge preview jika aktif
    if st.session_state.get('show_preview', False):
        st.markdown('<span class="preview-badge">👁️ MODE PREVIEW</span>', unsafe_allow_html=True)
        st.warning("⚠️ Anda sedang dalam mode preview. Klik 'Cetak Sekarang' untuk menghasilkan dokumen final.")
    
    # Folder template
    paths = get_default_paths()
    TND_FOLDER = paths['templates']
    
    # 🌟 REVISI 1: Otomatis buat folder jika belum ada (mencegah error OS)
    os.makedirs(TND_FOLDER, exist_ok=True)
    
    # Ambil daftar template
    all_templates = [f for f in os.listdir(TND_FOLDER) if f.endswith('.docx') and not f.startswith('~$')]
    
    # 🌟 REVISI 2: Validasi dinamis jika folder kosong atau tidak ada file .docx
    if not all_templates:
        st.warning(f"⚠️ Tidak ditemukan template `.docx` di folder target:")
        st.code(TND_FOLDER, language="text") # Menampilkan path asli (Windows/Linux) dengan rapi
        st.info("💡 Silakan masukkan file template surat Anda ke dalam folder tersebut lalu refresh halaman.")
        st.stop()
    
    col_select, col_info = st.columns([2, 1])
    with col_select:
        pilihan_template = st.selectbox("📂 Pilih Template Surat:", all_templates)
        template_terpilih_path = os.path.join(TND_FOLDER, pilihan_template)
    
    with col_info:
        st.info(f"📄 Template: {pilihan_template}")
        if st.button("🔄 Refresh", use_container_width=True):
            st.rerun()
    
    # Informasi template
    with st.expander("ℹ️ Informasi Template", expanded=False):
        st.markdown(f"""
        **Nama Template:** {pilihan_template}  
        **Lokasi:** {template_terpilih_path}  
        **Ukuran:** {os.path.getsize(template_terpilih_path) / 1024:.1f} KB  
        **Terakhir diubah:** {datetime.fromtimestamp(os.path.getmtime(template_terpilih_path)).strftime('%d %B %Y %H:%M')}
        """)
    
    st.markdown("---")
    
    # Tabs utama - disabled jika preview aktif
    tab_disabled = st.session_state.get('show_preview', False)
    
    if tab_disabled:
        st.info("🔒 Untuk mengedit data, tutup preview terlebih dahulu.")
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "✍️ Teks & Isi", 
        "📊 Tabel Data", 
        "🖼️ Gambar", 
        "👥 Penandatangan",
        "⚙️ Pengaturan"
    ])
    
    # ============================================
    # TAB 1: TEKS & ISI
    # ============================================
    with tab1:
        st.subheader("📝 Informasi Kepala Surat")
        
        with st.expander("🏷️ Identitas Surat", expanded=True):
            col_a, col_b = st.columns(2)
            
            with col_a:
                tahun = st.selectbox("Tahun", [2024, 2025, 2026, 2027], index=1, disabled=tab_disabled)
                bulan = st.selectbox("Bulan", [
                    (1, 'Januari'), (2, 'Februari'), (3, 'Maret'), (4, 'April'),
                    (5, 'Mei'), (6, 'Juni'), (7, 'Juli'), (8, 'Agustus'),
                    (9, 'September'), (10, 'Oktober'), (11, 'November'), (12, 'Desember')
                ], format_func=lambda x: x[1], index=0, disabled=tab_disabled)
                no_urut = st.number_input("Nomor Urut", min_value=1, max_value=999, value=1, disabled=tab_disabled)
                kode_unit = st.text_input("Kode Unit", "BP2JK-KT", disabled=tab_disabled)
                
                nomor_surat_auto = TNDEStandard.format_nomor_surat(no_urut, bulan[0], tahun, kode_unit)
                nomor_surat = st.text_input("Nomor Surat", nomor_surat_auto, disabled=tab_disabled)
                
                if nomor_surat and not TNDEStandard.validate_nomor_surat(nomor_surat):
                    st.warning("⚠️ Format nomor surat tidak sesuai TNDE. Format yang benar: 001/BP2JK-KT/01/2026")
            
            with col_b:
                hal = st.text_input("Perihal / Hal Surat", placeholder="Surat Izin Pelaksanaan Kegiatan", disabled=tab_disabled)
                lampiran = st.number_input("Jumlah Lampiran", min_value=0, max_value=99, value=1, disabled=tab_disabled)
                
                jenis_surat = st.selectbox(
                    "Jenis Surat",
                    ["Surat Izin", "Surat Tugas", "Surat Undangan", "Surat Keterangan", "Surat Keputusan", "Lainnya"],
                    disabled=tab_disabled
                )
                
                sifat_surat = st.radio(
                    "Sifat Surat",
                    ["Biasa", "Penting", "Segera", "Rahasia"],
                    horizontal=True,
                    disabled=tab_disabled
                )
        
        with st.expander("📄 Alinea Surat", expanded=True):
            template_alinea = st.selectbox(
                "Template Alinea Cepat",
                ["Kosong", "Surat Izin", "Surat Tugas", "Surat Undangan", "Surat Keterangan"],
                disabled=tab_disabled
            )
            
            templates = {
                "Surat Izin": {
                    "pembuka": "Menunjuk permohonan izin yang diajukan, dengan ini memberikan izin kepada:",
                    "isi": "Untuk melaksanakan kegiatan dengan jadwal dan pembagian waktu sebagai berikut:",
                    "penutup": "Demikian surat izin ini dibuat untuk dapat dipergunakan sebagaimana mestinya."
                },
                "Surat Tugas": {
                    "pembuka": "Berdasarkan Surat Keputusan Nomor ... tentang ... dengan ini menugaskan:",
                    "isi": "Untuk melaksanakan tugas sebagai berikut:",
                    "penutup": "Demikian surat tugas ini dibuat untuk dilaksanakan dengan sebaik-baiknya."
                },
                "Surat Undangan": {
                    "pembuka": "Dengan hormat, kami mengundang Bapak/Ibu untuk menghadiri acara:",
                    "isi": "Acara akan diselenggarakan pada:",
                    "penutup": "Demikian undangan ini kami sampaikan, atas perhatiannya kami ucapkan terima kasih."
                },
                "Surat Keterangan": {
                    "pembuka": "Yang bertanda tangan di bawah ini menerangkan bahwa:",
                    "isi": "Nama tersebut di atas adalah benar ...",
                    "penutup": "Demikian surat keterangan ini dibuat untuk dipergunakan sebagaimana mestinya."
                }
            }
            
            if template_alinea in templates:
                tpl = templates[template_alinea]
                pembuka_def = tpl["pembuka"]
                isi_def = tpl["isi"]
                penutup_def = tpl["penutup"]
            else:
                pembuka_def = ""
                isi_def = ""
                penutup_def = ""
            
            col_isi1, col_isi2 = st.columns(2)
            with col_isi1:
                txt_pembuka = st.text_area("Alinea Pembuka", pembuka_def, height=120, disabled=tab_disabled)
                txt_tempat = st.text_input("Tempat Surat", "Palangka Raya", disabled=tab_disabled)
                
            with col_isi2:
                txt_isi = st.text_area("Alinea Isi / Detail", isi_def, height=120, disabled=tab_disabled)
                tanggal_surat = st.date_input("Tanggal Surat", datetime.now(), disabled=tab_disabled)
        
        with st.expander("📝 Alinea Penutup & Pelengkap", expanded=False):
            txt_penutup = st.text_area("Alinea Penutup", penutup_def, height=100, disabled=tab_disabled)
            
            col_c1, col_c2 = st.columns(2)
            with col_c1:
                tembusan = st.text_area("Tembusan (pisahkan dengan koma)", 
                                       "Kepala BP2JK\nDirektur ...", disabled=tab_disabled)
            with col_c2:
                st.markdown("""
                **📌 Tips Penggunaan Placeholder:**
                - `[NOMOR_SURAT]` - untuk nomor surat
                - `{pembuka}` - untuk alinea pembuka
                - `$NAMA_ATASAN$` - untuk nama atasan
                - `(Alinea isi)` - untuk alinea isi
                
                **💡 Format akan dipertahankan** (Bold/Italic/Color)
                **⚡ Optimasi Split-Run:** Placeholder akan digabungkan otomatis
                """)
    
    # ============================================
    # TAB 2: TABEL DATA
    # ============================================
    with tab2:
        st.subheader("📊 Data Tabel (In-Body / Lampiran)")
        
        st.info("💡 **Tips Pro:** Jika nama kolom di bawah ini cocok dengan tabel di file Word, data akan otomatis mengisi tabel tersebut!")
        st.info("🔑 **Perbaikan:** Data akan dimapping berdasarkan **nama kolom**, bukan urutan. Jadi urutan kolom di Word tidak masalah!")
        
        col_kolom, col_opsi = st.columns([3, 1])
        with col_kolom:
            kolom_custom = st.text_input(
                "📋 Tentukan Nama Kolom (Pisahkan dengan koma):", 
                "NO, Jam, Hari Ke-1, Hari Ke-2, Hari Ke-3, Hari Ke-4",
                disabled=tab_disabled
            )
        
        with col_opsi:
            st.write("")
            st.write("")
            if st.button("🔄 Reset Kolom", disabled=tab_disabled):
                st.rerun()
        
        list_kolom = [k.strip() for k in kolom_custom.split(",") if k.strip() != ""]
        
        if list_kolom:
            data_contoh = {k: "" for k in list_kolom}
            
            sample_data = {
                "NO": "1",
                "Jam": "08.00-09.00",
                "Hari Ke-1": "Pengarahan",
                "Hari Ke-2": "Lanjutan",
                "Hari Ke-3": "Pelaksanaan",
                "Hari Ke-4": "Evaluasi",
                "Nama": "Contoh Nama",
                "Kegiatan": "Contoh Kegiatan"
            }
            
            for key, value in sample_data.items():
                if key in data_contoh:
                    data_contoh[key] = value
            
            df_dinamis = pd.DataFrame([data_contoh], columns=list_kolom)
        else:
            df_dinamis = pd.DataFrame()
            st.warning("⚠️ Silakan masukkan minimal 1 nama kolom")
        
        tabel_input = st.data_editor(
            df_dinamis, 
            num_rows="dynamic", 
            use_container_width=True,
            disabled=tab_disabled,
            column_config={
                "NO": st.column_config.NumberColumn("No", min_value=1, step=1),
            }
        )
        
        if not tabel_input.empty:
            st.caption(f"📊 Tabel berisi {len(tabel_input)} baris data dengan {len(tabel_input.columns)} kolom")
            
            if st.button("📥 Export ke CSV", disabled=tab_disabled):
                csv = tabel_input.to_csv(index=False)
                st.download_button(
                    label="Download CSV",
                    data=csv,
                    file_name=f"data_tabel_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv"
                )
    
    # ============================================
    # TAB 3: GAMBAR
    # ============================================
    with tab3:
        st.subheader("🖼️ Lampiran File Gambar")
        
        # Menggunakan key dinamis untuk reset
        gambar_upload = st.file_uploader(
            "Upload foto pendukung (PNG/JPG):", 
            type=["png", "jpg", "jpeg"],
            help="Ukuran gambar akan otomatis disesuaikan",
            disabled=tab_disabled,
            key=st.session_state.get("gambar_key", "uploader_awal")
        )
        
        if gambar_upload is not None:
            col_img1, col_img2 = st.columns(2)
            with col_img1:
                st.image(gambar_upload, caption="Preview Gambar", use_container_width=True)
            with col_img2:
                st.info(f"""
                **Informasi Gambar:**
                - Nama: {gambar_upload.name}
                - Ukuran: {gambar_upload.size / 1024:.1f} KB
                - Tipe: {gambar_upload.type}
                """)
                
                if st.button("🗑️ Hapus Gambar", disabled=tab_disabled):
                    st.session_state["gambar_key"] = f"uploader_{datetime.now().timestamp()}"
                    st.rerun()
    
    # ============================================
    # TAB 4: PENANDATANGAN
    # ============================================
    with tab4:
        st.subheader("👥 Konfigurasi Pejabat Penandatangan")
        
        col_ttd1, col_ttd2 = st.columns(2)
        with col_ttd1:
            jumlah_ttd = st.radio("Jumlah Penandatangan:", [1, 2, 3], horizontal=True, index=0, disabled=tab_disabled)
        
        with col_ttd2:
            tipe_ttd = st.selectbox(
                "Tipe Penandatanganan:",
                ["Pejabat", "Kuasa", "Pelaksana Tugas", "Pejabat Sementara"],
                disabled=tab_disabled
            )
        
        list_pejabat = []
        
        default_pejabat = [
            {
                "jabatan": "Kepala Balai Pelaksana Pemilihan Jasa Konstruksi Wilayah Kalimantan Tengah",
                "nama": "Nanang Rianto, S.Ant., M.PP.",
                "nip": "198510252008121001"
            },
            {
                "jabatan": "Pejabat Pembuat Komitmen",
                "nama": "Nama Pejabat Kedua",
                "nip": "19950812 202409 1 001"
            },
            {
                "jabatan": "Kepala Subbagian Tata Usaha",
                "nama": "Nama Pejabat Ketiga",
                "nip": "19871212 200801 2 002"
            }
        ]
        
        for i in range(jumlah_ttd):
            st.markdown(f"---")
            st.markdown(f"**Pejabat Ke-{i+1}:**")
            
            col_p1, col_p2, col_p3 = st.columns(3)
            
            with col_p1:
                jab = st.text_input(
                    f"Jabatan", 
                    default_pejabat[i]["jabatan"] if i < len(default_pejabat) else "",
                    key=f"jab_{i}",
                    disabled=tab_disabled
                )
            
            with col_p2:
                nam = st.text_input(
                    f"Nama Lengkap", 
                    default_pejabat[i]["nama"] if i < len(default_pejabat) else "",
                    key=f"nam_{i}",
                    disabled=tab_disabled
                )
            
            with col_p3:
                np = st.text_input(
                    f"NIP", 
                    default_pejabat[i]["nip"] if i < len(default_pejabat) else "",
                    key=f"nip_{i}",
                    disabled=tab_disabled
                )
            
            list_pejabat.append({"jabatan": jab, "nama": nam, "nip": np})
    
    # ============================================
    # TAB 5: PENGATURAN
    # ============================================
    with tab5:
        st.subheader("⚙️ Pengaturan Lanjutan")
        
        with st.expander("📁 Pengaturan Folder", expanded=True):
            col_f1, col_f2 = st.columns(2)
            with col_f1:
                folder_template = st.text_input("Folder Template", "C:\\doklib\\tnd", disabled=tab_disabled)
                folder_output = st.text_input("Folder Output", "C:\\doklib", disabled=tab_disabled)
            
            with col_f2:
                st.caption("📌 **Catatan:**")
                st.caption("• Template harus berformat .docx")
                st.caption("• Pastikan folder output memiliki akses tulis")
        
        with st.expander("🔧 Opsi Dokumen", expanded=True):
            col_o1, col_o2 = st.columns(2)
            with col_o1:
                auto_open = st.checkbox("Buka file setelah selesai", value=False, disabled=tab_disabled)
                preview_on_load = st.checkbox("Tampilkan preview sebelum cetak", value=True, disabled=tab_disabled)
            
            with col_o2:
                include_metadata = st.checkbox("Tambahkan metadata dokumen", value=True, disabled=tab_disabled)
                compress_images = st.checkbox("Kompres gambar", value=True, disabled=tab_disabled)
        
        with st.expander("📊 Log & Monitoring", expanded=False):
            if st.button("📋 Tampilkan Log Hari Ini", disabled=tab_disabled):
                log_file = os.path.join(r"C:\doklib\logs", f"apso_{datetime.now().strftime('%Y%m%d')}.log")
                if os.path.exists(log_file):
                    with open(log_file, 'r', encoding='utf-8') as f:
                        log_content = f.read()
                    st.text_area("Log Content", log_content, height=300)
                else:
                    st.info("Belum ada log untuk hari ini")
    
    # ============================================
    # 8. TOMBOL AKSI
    # ============================================
    st.markdown("---")
    
    # Jika preview aktif, tampilkan dialog preview
    if st.session_state.get('show_preview', False):
        if st.session_state.get('preview_buffer'):
            show_preview_dialog(
                st.session_state['preview_buffer'],
                st.session_state.get('preview_data'),
                pilihan_template
            )
        else:
            st.warning("⚠️ Preview tidak tersedia. Silakan generate ulang.")
            if st.button("🔄 Generate Ulang Preview"):
                st.session_state['show_preview'] = False
                st.rerun()
    else:
        # Tombol aksi normal
        col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
        with col_btn1:
            preview_button = st.button("👁️ Preview Dokumen", use_container_width=True)
        with col_btn2:
            proses_button = st.button("🚀 Cetak Dokumen", type="primary", use_container_width=True)
        with col_btn3:
            reset_button = st.button("🔄 Reset Semua", use_container_width=True)
        
        # Handle Preview
        if preview_button:
            try:
                with st.spinner("⏳ Membuat preview dokumen..."):
                    # Siapkan data
                    data_teks_gabungan = {
                        "[NOMOR_SURAT]": nomor_surat,
                        "NOMOR .../.../.../.../...": f"NOMOR {nomor_surat}",
                        "…./…./…./…./….": nomor_surat,
                        "Hal : ……………………………………………….": f"Hal : {hal}",
                        "Hal :": f"Hal : {hal}",
                        "pembuka": txt_pembuka,
                        "isi": txt_isi,
                        "penutup": txt_penutup,
                        "pembuka_isi": f"{txt_pembuka}\n\n{txt_isi}",
                        "tempat_tanggal": f"{txt_tempat}, {tanggal_surat.strftime('%d %B %Y')}",
                        "lampiran": str(lampiran),
                        "sifat": sifat_surat,
                        "jenis": jenis_surat,
                    }
                    
                    if list_pejabat:
                        data_teks_gabungan["[NAMA_ATASAN]"] = list_pejabat[0]["nama"]
                        data_teks_gabungan["[NIP_ATASAN]"] = list_pejabat[0]["nip"]
                        data_teks_gabungan["[JABATAN_ATASAN]"] = list_pejabat[0]["jabatan"]
                    
                    # Generate preview
                    preview_buffer = generate_preview_document(
                        template_terpilih_path,
                        data_teks_gabungan,
                        tabel_input,
                        gambar_upload,
                        list_pejabat
                    )
                    
                    # Simpan ke session state
                    st.session_state['preview_buffer'] = preview_buffer
                    st.session_state['show_preview'] = True
                    
                    # Data ringkasan
                    st.session_state['preview_data'] = {
                        'template': pilihan_template,
                        'nomor_surat': nomor_surat,
                        'jumlah_ttd': len(list_pejabat),
                        'baris_tabel': len(tabel_input) if not tabel_input.empty else 0,
                        'ada_gambar': gambar_upload is not None,
                        'alinea_count': sum(1 for x in [txt_pembuka, txt_isi, txt_penutup] if x),
                        'hal': hal,
                        'tempat': txt_tempat,
                        'tanggal': tanggal_surat.strftime('%d %B %Y'),
                        'jenis_surat': jenis_surat,
                        'sifat_surat': sifat_surat,
                    }
                    
                    st.rerun()
                    
            except Exception as e:
                st.error(f"❌ Gagal membuat preview: {e}")
                logger.error(f"Preview error: {e}")
        
        # Handle Cetak
        if proses_button or st.session_state.get('confirm_print', False):
            if st.session_state.get('confirm_print', False):
                st.session_state['confirm_print'] = False
            
            # Siapkan data
            data_teks_gabungan = {
                "[NOMOR_SURAT]": nomor_surat,
                "NOMOR .../.../.../.../...": f"NOMOR {nomor_surat}",
                "…./…./…./…./….": nomor_surat,
                "Hal : ……………………………………………….": f"Hal : {hal}",
                "Hal :": f"Hal : {hal}",
                "pembuka": txt_pembuka,
                "isi": txt_isi,
                "penutup": txt_penutup,
                "pembuka_isi": f"{txt_pembuka}\n\n{txt_isi}",
                "tempat_tanggal": f"{txt_tempat}, {tanggal_surat.strftime('%d %B %Y')}",
                "lampiran": str(lampiran),
                "sifat": sifat_surat,
                "jenis": jenis_surat,
            }
            
            if list_pejabat:
                data_teks_gabungan["[NAMA_ATASAN]"] = list_pejabat[0]["nama"]
                data_teks_gabungan["[NIP_ATASAN]"] = list_pejabat[0]["nip"]
                data_teks_gabungan["[JABATAN_ATASAN]"] = list_pejabat[0]["jabatan"]
            
            # Validasi
            errors, warnings = validate_inputs(data_teks_gabungan, tabel_input, list_pejabat)
            
            if errors:
                for error in errors:
                    st.error(error)
                st.stop()
            
            if warnings:
                for warning in warnings:
                    st.warning(warning)
            
            # Proses dokumen
            output_file_path = os.path.join(r"C:\doklib", f"Hasil_Final_{pilihan_template}")
            
            metadata = {
                'judul': f"Surat {jenis_surat} - {hal}",
                'perihal': hal,
                'pembuat': 'APSO Ultimate - BP2JK Kalteng',
                'tanggal': datetime.now().isoformat()
            }
            
            try:
                with st.spinner('🔄 Memproses dokumen... Mohon tunggu'):
                    success = proses_dokumen_universal(
                        template_terpilih_path,
                        output_file_path,
                        data_teks_gabungan,
                        tabel_input,
                        gambar_upload,
                        list_pejabat,
                        metadata
                    )
                
                if success:
                    st.success(f"✅ Dokumen berhasil dibuat!")
                    
                    file_size = os.path.getsize(output_file_path) / 1024
                    st.info(f"📄 File size: {file_size:.1f} KB | Lokasi: {output_file_path}")
                    
                    # Tampilkan 3 opsi: Download, Buka, Preview
                    col_d1, col_d2, col_d3 = st.columns(3)
                    
                    with col_d1:
                        with open(output_file_path, "rb") as file:
                            st.download_button(
                                label="📥 Download Dokumen",
                                data=file,
                                file_name=f"Surat_{nomor_surat}_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    
                    with col_d2:
                        if os.name == 'nt':
                            if st.button("📂 Buka di Word", use_container_width=True):
                                try:
                                    os.startfile(output_file_path)
                                except Exception as e:
                                    st.error(f"Gagal membuka dokumen: {e}")
                        else:
                            st.info("💡 Fitur 'Buka di Word' hanya tersedia di Windows")
                    
                    with col_d3:
                        if st.button("👁️ Lihat Preview", use_container_width=True):
                            try:
                                with open(output_file_path, "rb") as f:
                                    preview_buffer = io.BytesIO(f.read())
                                    st.session_state['preview_buffer'] = preview_buffer
                                    st.session_state['show_preview'] = True
                                    st.rerun()
                            except Exception as e:
                                st.error(f"Gagal membuka preview: {e}")
                    
                    if auto_open and os.name == 'nt':
                        try:
                            os.startfile(output_file_path)
                        except Exception as e:
                            logger.warning(f"Auto open failed: {e}")
                    
            except Exception as e:
                st.error(f"❌ Terjadi kesalahan: {str(e)}")
                logger.error(f"Error: {str(e)}")
                
                with st.expander("🔍 Detail Error"):
                    st.code(str(e), language="python")
        
        # Handle Reset
        if reset_button:
            for key in ['show_preview', 'confirm_print', 'preview_buffer', 'preview_data']:
                if key in st.session_state:
                    st.session_state[key] = None if key != 'show_preview' else False
            st.session_state['gambar_key'] = f"uploader_{datetime.now().timestamp()}"
            st.rerun()

# ============================================
# 9. MAIN EXECUTION
# ============================================

if __name__ == "__main__":
    # Tampilkan info versi saat startup
    app_info = get_app_info()
    logger.info(f"🚀 {app_info['name']} v{app_info['version']} dimulai")
    logger.info(f"📝 TNDE Compliant: {app_info['tnde_compliant']}")
    main()